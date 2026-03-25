import pdfplumber
import pandas as pd
from docx import Document
import re
import json
import asyncio
from datetime import datetime
from PIL import Image
import io
import base64
from config import config, INVOICE_EXTRACTION_PROMPT, BATCH_EXTRACTION_PROMPT
from typing import Dict, List, Optional, Union


class InvoiceExtractor:
    """Enhanced invoice extractor with Gemini AI integration"""
    
    def __init__(self):
        self.config = config
        self.supported_image_formats = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp'}
        
    async def extract_with_gemini(self, content: Union[str, Image.Image], is_image: bool = False) -> Optional[Dict]:
        """Extract invoice data using Gemini AI"""
        if not self.config.is_available():
            return None
            
        try:
            if is_image:
                # Process image with Gemini Vision
                prompt = INVOICE_EXTRACTION_PROMPT
                response = await self._generate_content_async(prompt, content)
            else:
                # Process text with Gemini
                prompt = INVOICE_EXTRACTION_PROMPT + content
                response = await self._generate_content_async(prompt)
            
            if response:
                # Parse JSON response
                json_text = self._extract_json_from_response(response)
                if json_text:
                    data = json.loads(json_text)
                    return self._normalize_extracted_data(data)
                    
        except Exception as e:
            print(f"Gemini extraction error: {e}")
            return None
    
    async def _generate_content_async(self, prompt: str, image: Optional[Image.Image] = None) -> Optional[str]:
        """Async wrapper for Gemini content generation"""
        try:
            if image:
                # Use Gemini Vision for image + text
                response = self.config.model.generate_content([prompt, image])
            else:
                # Use Gemini for text only
                response = self.config.model.generate_content(prompt)
            
            if response and response.text:
                return response.text.strip()
        except Exception as e:
            error_text = str(e)
            print(f"Gemini API error: {error_text}")
            error_lower = error_text.lower()
            if (
                "403" in error_text
                or "api key" in error_lower
                or "leaked" in error_lower
                or "permission" in error_lower
                or "unauthorized" in error_lower
            ):
                self.config.disable_ai(error_text)
            return None
    
    def _extract_json_from_response(self, response: str) -> Optional[str]:
        """Extract JSON from Gemini response"""
        # Look for JSON blocks in response
        patterns = [
            r'```json\s*(\{.*?\})\s*```',
            r'```\s*(\{.*?\})\s*```', 
            r'(\{[^}]*"customer"[^}]*\})',
            r'(\[[^\]]*\{[^}]*"customer"[^}]*\}[^\]]*\])'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, response, re.DOTALL | re.IGNORECASE)
            if match:
                return match.group(1)
        
        # If no JSON block found, try to extract JSON-like content
        if '{' in response and '}' in response:
            start = response.find('{')
            end = response.rfind('}') + 1
            return response[start:end]
            
        return None
    
    def _normalize_extracted_data(self, data: Union[Dict, List]) -> Dict:
        """Normalize extracted data to consistent format"""
        if isinstance(data, list) and data:
            # Take first invoice if multiple found
            data = data[0]
        
        if not isinstance(data, dict):
            return {"customer": "", "date": "", "item": "", "amount": ""}
        
        # Normalize field names and clean data
        normalized = {
            "customer": str(data.get("customer", "")).strip(),
            "date": str(data.get("date", "")).strip(),
            "item": str(data.get("item", "")).strip(), 
            "amount": str(data.get("amount", "")).strip(),
            "confidence": data.get("confidence", "Medium"),
            "invoice_number": data.get("invoice_number", ""),
            "currency": data.get("currency", "")
        }
        
        # Clean amount field - keep only numbers and decimal points
        if normalized["amount"]:
            amount_clean = re.sub(r'[^\d.,]', '', normalized["amount"])
            normalized["amount"] = amount_clean
        
        return normalized

    def extract_pdf_traditional(self, file_path: str) -> Optional[Dict]:
        """Traditional PDF extraction fallback"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"Error reading PDF: {e}")
            return None

        return self.parse_text_traditional(text)

    def extract_excel_traditional(self, file_path: str) -> Optional[List[Dict]]:
        """Traditional Excel extraction"""
        try:
            df = pd.read_excel(file_path)
            data = []

            for _, row in df.iterrows():
                invoice_data = {
                    "customer": str(row.get("Customer", row.get("customer", ""))).strip(),
                    "date": str(row.get("Date", row.get("date", ""))).strip(),
                    "item": str(row.get("Item", row.get("item", ""))).strip(),
                    "amount": str(row.get("Amount", row.get("amount", ""))).strip()
                }
                
                if any(invoice_data.values()):
                    data.append(invoice_data)

            return data if data else None
        
        except Exception as e:
            print(f"Error reading Excel: {e}")
            return None

    def extract_word_traditional(self, file_path: str) -> Optional[Dict]:
        """Traditional Word extraction fallback"""
        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return self.parse_text_traditional(text)
        
        except Exception as e:
            print(f"Error reading Word document: {e}")
            return None

    def parse_text_traditional(self, text: str) -> Dict:
        """Traditional regex-based text parsing"""
        data = {
            "customer": "",
            "date": "",
            "item": "",
            "amount": "",
            "confidence": "Low"
        }

        if not text:
            return data

        lines = text.split("\n")
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            line_lower = line.lower()
            
            # Extract customer information
            customer_patterns = [
                r"customer[:\s]*(.+)",
                r"bill to[:\s]*(.+)",
                r"client[:\s]*(.+)",
                r"company[:\s]*(.+)"
            ]
            
            for pattern in customer_patterns:
                match = re.search(pattern, line_lower)
                if match and not data["customer"]:
                    data["customer"] = match.group(1).strip()
                    break

            # Extract date information
            date_patterns = [
                r"date[:\s]*(.+)",
                r"invoice date[:\s]*(.+)",
                r"issued[:\s]*(.+)"
            ]
            
            for pattern in date_patterns:
                match = re.search(pattern, line_lower)
                if match and not data["date"]:
                    date_str = match.group(1).strip()
                    data["date"] = date_str[:50]
                    break

            # Extract item information
            item_patterns = [
                r"item[:\s]*(.+)",
                r"product[:\s]*(.+)",
                r"description[:\s]*(.+)",
                r"service[:\s]*(.+)"
            ]
            
            for pattern in item_patterns:
                match = re.search(pattern, line_lower)
                if match and not data["item"]:
                    data["item"] = match.group(1).strip()
                    break

            # Extract amount information
            amount_patterns = [
                r"total[:\s]*(.+)",
                r"amount[:\s]*(.+)",
                r"sum[:\s]*(.+)",
                r"price[:\s]*(.+)"
            ]
            
            for pattern in amount_patterns:
                match = re.search(pattern, line_lower)
                if match and not data["amount"]:
                    amount_str = match.group(1).strip()
                    amount_match = re.search(r'[\d,]+\.?\d*', amount_str)
                    if amount_match:
                        data["amount"] = amount_match.group(0)
                    break

        return data

    def process_image(self, file_path: str) -> Optional[Image.Image]:
        """Process image file for Gemini vision"""
        try:
            image = Image.open(file_path)
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Resize if too large (Gemini has size limits)
            max_size = (1024, 1024)
            if image.size[0] > max_size[0] or image.size[1] > max_size[1]:
                image.thumbnail(max_size, Image.Resampling.LANCZOS)
            
            return image
        except Exception as e:
            print(f"Error processing image: {e}")
            return None

# Global extractor instance
extractor = InvoiceExtractor()

async def extract_invoice(file_path: str) -> Optional[Union[Dict, List[Dict]]]:
    """
    Main async function to extract invoice data from various file types
    """
    if not file_path:
        return None
    
    file_path_lower = file_path.lower()
    
    try:
        # Determine file type and processing method
        if any(file_path_lower.endswith(ext) for ext in extractor.supported_image_formats):
            # Process image files
            return await extract_image_invoice(file_path)
            
        elif file_path_lower.endswith(".pdf"):
            return await extract_pdf_invoice(file_path)
        
        elif file_path_lower.endswith((".xlsx", ".xls")):
            return await extract_excel_invoice(file_path)
            
        elif file_path_lower.endswith(".csv"):
            return await extract_csv_invoice(file_path)
        
        elif file_path_lower.endswith(".docx"):
            return await extract_word_invoice(file_path)
            
        elif file_path_lower.endswith(".txt"):
            return await extract_text_invoice(file_path)
        
        else:
            print(f"Unsupported file type: {file_path}")
            return None
    
    except Exception as e:
        print(f"Error processing file {file_path}: {e}")
        return None

async def extract_image_invoice(file_path: str) -> Optional[Dict]:
    """Extract data from image files using Gemini Vision"""
    # Process image
    image = extractor.process_image(file_path)
    if not image:
        return None
    
    # Try Gemini Vision first
    if extractor.config.is_available():
        result = await extractor.extract_with_gemini(image, is_image=True)
        if result and any(result.values()):
            print("✅ Extracted with Gemini Vision")
            return result
    
    print("⚠️  Image processing requires Gemini API")
    return {"customer": "", "date": "", "item": "", "amount": "", "confidence": "Failed - No OCR"}

async def extract_pdf_invoice(file_path: str) -> Optional[Dict]:
    """Extract data from PDF files"""
    # Extract text first
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return None
    
    # Try Gemini extraction first
    if extractor.config.is_available() and text.strip():
        result = await extractor.extract_with_gemini(text)
        if result and any(result.values()):
            print("✅ Extracted with Gemini AI")
            return result
    
    # Fallback to traditional method 
    if extractor.config.fallback_to_traditional:
        result = extractor.parse_text_traditional(text)
        if result and any(result.values()):
            print("✅ Extracted with traditional parsing")
            return result
    
    return {"customer": "", "date": "", "item": "", "amount": "", "confidence": "Low"}

async def extract_excel_invoice(file_path: str) -> Optional[Union[Dict, List[Dict]]]:
    """Extract data from Excel files"""
    result = extractor.extract_excel_traditional(file_path)
    if result:
        print("✅ Extracted from Excel")
        return result[0] if len(result) == 1 else result
    return None

async def extract_word_invoice(file_path: str) -> Optional[Dict]:
    """Extract data from Word files"""
    # Extract text
    try:
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        print(f"Error reading Word document: {e}")
        return None
    
    # Try Gemini first
    if extractor.config.is_available() and text.strip():
        result = await extractor.extract_with_gemini(text)
        if result and any(result.values()):
            print("✅ Extracted with Gemini AI")
            return result
    
    # Fallback to traditional
    if extractor.config.fallback_to_traditional:
        result = extractor.parse_text_traditional(text)
        if result and any(result.values()):
            print("✅ Extracted with traditional parsing") 
            return result
    
    return {"customer": "", "date": "", "item": "", "amount": "", "confidence": "Low"}

async def extract_csv_invoice(file_path: str) -> Optional[Union[Dict, List[Dict]]]:
    """Extract data from CSV files"""
    try:
        df = pd.read_csv(file_path)
        print(f"CSV Debug: DataFrame shape: {df.shape}")
        print(f"CSV Debug: DataFrame columns: {list(df.columns)}")
        print(f"CSV Debug: DataFrame head:\n{df.head()}")
        
        data = []

        for i, row in df.iterrows():
            invoice_data = {
                "customer": str(row.get("Customer", row.get("customer", ""))).strip(),
                "date": str(row.get("Date", row.get("date", ""))).strip(),
                "item": str(row.get("Item", row.get("item", ""))).strip(),
                "amount": str(row.get("Amount", row.get("amount", ""))).strip(),
                "confidence": "High"
            }
            
            print(f"CSV Debug: Row {i}: {invoice_data}")
            
            if any(invoice_data.values()):
                data.append(invoice_data)

        print(f"CSV Debug: Total records processed: {len(data)}")
        
        if data:
            print(f"✅ Extracted {len(data)} records from CSV")
            return data  # Always return the full list, let app.py handle single vs multiple
        return None
        
    except Exception as e:
        print(f"Error reading CSV: {e}")
        return None

async def extract_text_invoice(file_path: str) -> Optional[Dict]:
    """Extract data from text files"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
    except Exception as e:
        print(f"Error reading text file: {e}")
        return None
    
    # Try Gemini first
    if extractor.config.is_available() and text.strip():
        result = await extractor.extract_with_gemini(text)
        if result and any(result.values()):
            print("✅ Extracted with Gemini AI")
            return result
    
    # Fallback to traditional
    if extractor.config.fallback_to_traditional:
        result = extractor.parse_text_traditional(text)
        if result and any(result.values()):
            print("✅ Extracted with traditional parsing") 
            return result
    
    return {"customer": "", "date": "", "item": "", "amount": "", "confidence": "Low"}